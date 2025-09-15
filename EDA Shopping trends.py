#import libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import io, sys
buffer = io.StringIO()
sys.stdout = buffer
data = pd.read_csv("C:\\Users\\Avaneesh\\Downloads\\archive (2)\\shopping_trends.csv")#To read the csv file
print(data.describe())#Shows the columns 
print("Top 5 Rows")
print(data.head())#Print the first 5 rows 
print("Bottom 5 Rows")
print(data.tail())#Print the last 5 rows 
data["Gender"] = data["Gender"].astype("category")#Convert the Gender datatype(Str) into category datatype
print(data.dtypes)
print("The size of the CSV dataset is: ",data.shape)#Returns the size of the data(rows x columns)
print("The total number of NULL values present in the datas by columns :")
print(data.isna().sum())#Find the number of missing value
avg=data['Purchase Amount (USD)'].mean()#Average of Purchased amount
print('Average of the Purchased Amount is ',avg)
print("The maximum amount purchased is: ",data["Purchase Amount (USD)"].max())#Maximum amnt purchased
print("The minimum amount purchased is: ",data["Purchase Amount (USD)"].min())#Minimum amnt Purchased
print("NO.of duplicates",data.duplicated().sum())#Displays the numb of duplicate values
print("No. of Colors purchased : " ,data["Color"].value_counts())#Total no. of each color
print(data["Payment Method"].value_counts())#Display the number of specified value
plt.figure()
data['Age'].plot(
    kind="hist",
    bins=15,          
    ylim=(0, 25),
) #Histogram plot for Univariate analysis
plt.title("Histogram Plot of Customer Age")
plt.xlabel("Ages")
plt.ylabel("Count")
plt.show()
plt.figure()
plt.scatter(data=data,y="Review Rating",x="Purchase Amount (USD)")#Scatter plot for Bivariate analysis
plt.title("Purchase amount VS Review rating")
plt.xlabel("Purchase amount")
plt.ylabel("Review Rating")
plt.show()
gp=data.groupby("Size")["Purchase Amount (USD)"].sum()
print(gp)
print("Total amount of purchased: ",data["Purchase Amount (USD)"].sum())#Calculating in specific column
loc=data.groupby("Location")["Shipping Type"].value_counts().head()#Used to group two columns 
print(loc)
plt.figure()
plt.scatter(data=data,x='Gender',y="Review Rating")#Scatter plot for Review rating
plt.xlabel("Gender")
plt.ylabel("Ratings")
plt.title("Ratings by Gender")
plt.show()
sys.stdout = sys.__stdout__
all_output = buffer.getvalue()#To get all captured outputs
doc = Document()#To create Word document and write outputs
doc.add_heading('Exploratory Data Analysis Report', level=1)#Adding headings
doc.add_heading('shopping Trend Analysis', level=2)
doc.add_paragraph(all_output)#Adding the outputs
doc.add_heading('Visualizations', level=2)
doc.add_paragraph('Histogram of Age')
doc.add_picture("C:\\Users\\Avaneesh\\Downloads\\hist_age.png", width=Inches(5))#Adding the visualization pictures
doc.add_paragraph('Purchase Amount vs Review Rating')
doc.add_picture("C:\\Users\\Avaneesh\\Downloads\\scatter_purchase_rating.png", width=Inches(5))
doc.add_paragraph('Gender vs Ratings')
doc.add_picture("C:\\Users\\Avaneesh\\Downloads\\scatter_gender_rating.png", width=Inches(5))
doc.save("C:\\Users\\Avaneesh\\Downloads\\EDA_Output_Reportnew.docx")#Saving as Document
print("Report saved at C:\\Users\\Avaneesh\\Downloads\\EDA_Output_Reportnew.docx")